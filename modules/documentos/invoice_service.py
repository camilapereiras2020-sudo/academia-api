"""
Invoice/Receipt generation service for Cami&Co Academia.
Clean version - no corrupted characters.
"""
import os
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.join(os.path.expanduser("~"), "Desktop", "CAMI&CO", "Facturas", "DE CAM&CO", "Facturas 2026", "Control Facturas")
WORD_DIR = os.path.join(BASE_DIR, "Word")
PDF_DIR = os.path.join(BASE_DIR, "PDFs")
EXCEL_DIR = os.path.join(BASE_DIR, "Excel")
BASE_INVOICE_DIR = os.path.join(os.path.expanduser("~"), "Desktop", "CAMI&CO", "Facturas", "DE CAM&CO")
LOGO_PATH = os.path.join(os.path.dirname(__file__), "logo.png")

GOLD = RGBColor(0xB0, 0x8D, 0x57)
DARK = RGBColor(0x2D, 0x2D, 0x2D)
GRAY = RGBColor(0x6B, 0x6B, 0x6B)
LIGHT_GRAY = RGBColor(0x95, 0x95, 0x95)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GOLD_HEX = "B08D57"
DARK_HEX = "2D2D2D"
LIGHT_BG = "F7F5F2"

IBAN = "ES10 2080 5020 0530 4003 9725"

MESES_ES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
}

METODOS_FACTURA = ["bizum", "transferencia", "tarjeta", "online"]


def es_factura(metodo):
    return metodo.lower() in METODOS_FACTURA if metodo else True


def tipo_fiscal(pagador_nif):
    return "F1" if pagador_nif and pagador_nif.strip() else "F2"


def _remove_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def _set_cell_shading(cell, color_hex):
    shading = parse_xml(
        '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' w:fill="{}" w:val="clear"/>'.format(color_hex)
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_borders(cell, color="DDDDDD", sz="2"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{c}"/>'
        '<w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{c}"/>'
        '<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{c}"/>'
        '<w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{c}"/>'
        '</w:tcBorders>'.format(sz=sz, c=color)
    )
    tcPr.append(tcBorders)


def _add_run(paragraph, text, size=10, color=None, bold=False, italic=False, font="Century Gothic"):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font
    if color:
        run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return run


def _get_trimester(month):
    if month <= 4:
        return "T1 (Enero-Abril)"
    elif month <= 8:
        return "T2 (Mayo-Agosto)"
    else:
        return "T3 (Septiembre-Diciembre)"


def _get_output_dir(year, month):
    trimester = _get_trimester(month)
    path = os.path.join(BASE_INVOICE_DIR, "Facturas {}".format(year), "PDFs", trimester)
    os.makedirs(path, exist_ok=True)
    return path


def _next_invoice_number(academia_id, prefix="CC"):
    year_short = str(date.today().year)[2:]
    year_full = date.today().year
    year_dir = os.path.join(BASE_INVOICE_DIR, "Facturas {}".format(year_full), "PDFs")

    max_num = 236
    if os.path.exists(year_dir):
        for root, dirs, files in os.walk(year_dir):
            for f in files:
                if f.upper().startswith(prefix.upper()) and (f.endswith(".pdf") or f.endswith(".docx")):
                    try:
                        name = f.replace(".pdf", "").replace(".docx", "")
                        num_part = name.upper().replace(prefix.upper(), "").replace("-{}".format(year_short), "")
                        num = int(num_part)
                        max_num = max(max_num, num)
                    except ValueError:
                        pass
    return "{}{}{}{}".format(prefix, max_num + 1, "-", year_short)


def _format_date_es(d):
    return "{} de {} de {}".format(d.day, MESES_ES[d.month], d.year)


def _format_eur(amount):
    return "{:,.2f} EUR".format(float(amount)).replace(",", "X").replace(".", ",").replace("X", ".")


def _add_border_line(doc, color_hex, sz="12", position="bottom"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:{pos} w:val="single" w:sz="{sz}" w:space="1" w:color="{c}"/>'
        '</w:pBdr>'.format(pos=position, sz=sz, c=color_hex)
    )
    pPr.append(pBdr)
    return p


def generate_invoice(
    academia_nombre="Camila Pereiras Casal",
    academia_dir="C/ Pedro Soto Couselo 5, 2 B",
    academia_ciudad="36995, Poio (Pontevedra)",
    academia_tel="698 183 419",
    academia_cif="39468659S",
    pagador_nombre="",
    pagador_nif="",
    pagador_dir="",
    pagador_ciudad="",
    pagador_tel="",
    pagador_email="",
    alumno_nombre="",
    grupo_nombre="",
    periodo="2026-04",
    mensualidad=0,
    descuento=0,
    extras=None,
    total=0,
    metodo="",
    concepto_libre="",
    num_doc=None,
    fecha=None,
    tipo="factura",
    academia_id=1,
):
    if extras is None:
        extras = []
    if fecha is None:
        fecha = date.today()
    if isinstance(fecha, str):
        fecha = date.fromisoformat(fecha)
    if num_doc is None:
        num_doc = _next_invoice_number(academia_id)

    year_str, month_str = periodo.split("-")
    year, month = int(year_str), int(month_str)
    mes_nombre = MESES_ES[month].capitalize()

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Century Gothic"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # Header table
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    for cell in header_table.columns[0].cells:
        cell.width = Cm(8)
    for cell in header_table.columns[1].cells:
        cell.width = Cm(9)

    left_cell = header_table.rows[0].cells[0]
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_logo = left_cell.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(LOGO_PATH):
        run = p_logo.add_run()
        run.add_picture(LOGO_PATH, width=Cm(4.5))
    else:
        _add_run(p_logo, "Cami&Co", size=24, color=GOLD, bold=True, font="Book Antiqua")

    right_cell = header_table.rows[0].cells[1]
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_text = "FACTURA" if tipo == "factura" else "RECIBO"
    _add_run(p, title_text, size=32, color=GOLD, bold=True, font="Book Antiqua")

    p = right_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(p, "N {} {}".format(chr(186), num_doc), size=11, color=DARK)

    p = right_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(p, "Fecha: {}".format(_format_date_es(fecha)), size=10, color=GRAY)

    for row in header_table.rows:
        for cell in row.cells:
            _remove_borders(cell)

    _add_border_line(doc, GOLD_HEX)

    # Quote - Mandela
    p_quote = doc.add_paragraph()
    p_quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_quote.paragraph_format.space_before = Pt(4)
    p_quote.paragraph_format.space_after = Pt(12)
    _add_run(
        p_quote,
        '"It always seems impossible until it\'s done."',
        size=8.5, color=LIGHT_GRAY, italic=True, font="Book Antiqua"
    )
    _add_run(p_quote, "  -- Nelson Mandela", size=8, color=LIGHT_GRAY, font="Book Antiqua")

    # DE / PARA
    info_table = doc.add_table(rows=1, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in info_table.columns[0].cells:
        cell.width = Cm(8.5)
    for cell in info_table.columns[1].cells:
        cell.width = Cm(8.5)

    emisor_cell = info_table.rows[0].cells[0]
    p = emisor_cell.paragraphs[0]
    _add_run(p, "DE:", size=8, color=GOLD, bold=True)
    for text, bold, size in [
        (academia_nombre, True, 11),
        (academia_dir, False, 9.5),
        (academia_ciudad, False, 9.5),
        ("Tel: {}".format(academia_tel), False, 9.5),
        ("CIF: {}".format(academia_cif), False, 9.5),
        ("IBAN: {}".format(IBAN), False, 9),
    ]:
        p = emisor_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        _add_run(p, text, size=size, color=DARK if bold else GRAY, bold=bold)

    cliente_cell = info_table.rows[0].cells[1]
    p = cliente_cell.paragraphs[0]
    _add_run(p, "PARA:", size=8, color=GOLD, bold=True)
    client_fields = [
        ("Nombre", pagador_nombre),
        ("DNI/NIF", pagador_nif),
        ("Direccion", pagador_dir),
        ("Ciudad / CP", pagador_ciudad),
        ("Telefono", pagador_tel or pagador_email),
    ]
    for label, value in client_fields:
        if not value:
            continue
        p = cliente_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        _add_run(p, "{}: ".format(label), size=9, color=LIGHT_GRAY)
        _add_run(p, value, size=9.5, color=DARK)

    for row in info_table.rows:
        for cell in row.cells:
            _remove_borders(cell)

    # Alumno
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    _add_run(p, "ALUMNO/S", size=8, color=GOLD, bold=True)

    alumno_table = doc.add_table(rows=1, cols=1)
    alumno_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = alumno_table.rows[0].cells[0]
    cell.width = Cm(17)
    _set_cell_shading(cell, LIGHT_BG)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    _add_run(p, alumno_nombre, size=11, color=DARK, bold=True)
    _set_cell_borders(cell, GOLD_HEX, "4")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)

    # Items
    items = []
    if concepto_libre:
        desc_main = concepto_libre
    else:
        desc_main = "Clases de inglés -- {} {}".format(mes_nombre, year)
        if grupo_nombre:
            desc_main += " ({})".format(grupo_nombre)

    items.append((desc_main, "1", _format_eur(mensualidad), _format_eur(mensualidad)))

    if float(descuento) > 0:
        items.append(("Descuento", "1", "-{}".format(_format_eur(descuento)), "-{}".format(_format_eur(descuento))))

    for ex in extras:
        concepto = ex.get("concepto", "Extra")
        importe = float(ex.get("importe", 0))
        items.append((concepto, "1", _format_eur(importe), _format_eur(importe)))

    num_item_rows = max(len(items), 1)
    num_rows = 1 + num_item_rows + 3
    items_table = doc.add_table(rows=num_rows, cols=4)
    items_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_widths = [Cm(9), Cm(2.5), Cm(2.75), Cm(2.75)]
    for row in items_table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = w

    headers = ["Descripcion", "Cantidad", "Precio", "Importe"]
    for ci, h in enumerate(headers):
        cell = items_table.rows[0].cells[ci]
        _set_cell_shading(cell, DARK_HEX)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if ci >= 2 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        _add_run(p, h, size=9, color=WHITE, bold=True)

    for ri, (desc, cant, precio, importe) in enumerate(items):
        row_idx = ri + 1
        for ci, val in enumerate([desc, cant, precio, importe]):
            cell = items_table.rows[row_idx].cells[ci]
            if ri % 2 == 0:
                _set_cell_shading(cell, LIGHT_BG)
            p = cell.paragraphs[0]
            if ci >= 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif ci == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            _add_run(p, val, size=10, color=DARK)

    extras_total = sum(float(ex.get("importe", 0)) for ex in extras)
    subtotal = float(mensualidad) - float(descuento)

    summary = [
        ("SUBTOTAL", _format_eur(subtotal), False),
        ("OTROS", _format_eur(extras_total), False),
        ("TOTAL", _format_eur(total), True),
    ]

    for si, (label, amount, is_total) in enumerate(summary):
        row_idx = num_item_rows + 1 + si
        row = items_table.rows[row_idx]
        cell_label = row.cells[0]
        cell_label.merge(row.cells[2])
        p = cell_label.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        _add_run(p, label, size=10, color=DARK, bold=True)

        cell_amt = row.cells[3]
        if is_total:
            _set_cell_shading(cell_amt, GOLD_HEX)
        p = cell_amt.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        _add_run(p, amount, size=11 if is_total else 10,
                 color=WHITE if is_total else DARK, bold=is_total)

    for row in items_table.rows:
        for cell in row.cells:
            _set_cell_borders(cell)

    if metodo:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        _add_run(p, "Metodo de pago: ", size=9, color=GRAY)
        _add_run(p, metodo.capitalize(), size=9, color=DARK, bold=True)

    # Footer
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="6" w:space="6" w:color="{}"/>'.format(GOLD_HEX) +
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    _add_run(p, "TÉRMINOS DE PAGO", size=8, color=DARK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, "Los honorarios se gestionaran dentro de los ", size=8, color=GRAY)
    _add_run(p, "primeros 5 dias naturales del mes", size=8, color=DARK, bold=True)
    _add_run(p, ".", size=8, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    _add_run(p, IBAN, size=8, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    _add_run(
        p,
        "Operacion exenta de IVA segun el art. 20.Uno de la Ley 37/1992",
        size=8, color=LIGHT_GRAY, italic=True
    )

    # Save
    output_dir = _get_output_dir(fecha.year, fecha.month)
    docx_filename = "{}.docx".format(num_doc)
    pdf_filename = "{}.pdf".format(num_doc)
    docx_path = os.path.join(output_dir, docx_filename)
    pdf_path = os.path.join(output_dir, pdf_filename)

    doc.save(docx_path)

    try:
        import pythoncom
        pythoncom.CoInitialize()
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        pythoncom.CoUninitialize()
        if not os.path.exists(pdf_path):
            pdf_path = None
    except Exception as e:
        print("PDF conversion error: {}".format(e))
        pdf_path = None

    return num_doc, pdf_path or docx_path, docx_path


def update_excel_log(
    num_doc, fecha, pagador_nombre, pagador_nif, pagador_dir,
    alumno_nombre, grupo_nombre, periodo, mensualidad, descuento,
    extras_total, total, metodo, estado, es_factura_fiscal,
):
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    if isinstance(fecha, str):
        fecha = date.fromisoformat(fecha)

    year = fecha.year
    excel_dir = os.path.join(BASE_INVOICE_DIR, "Facturas {}".format(year))
    os.makedirs(excel_dir, exist_ok=True)
    excel_path = os.path.join(excel_dir, "Registro_Ingresos_{}.xlsx".format(year))

    gold_fill = PatternFill(start_color="B08D57", end_color="B08D57", fill_type="solid")
    dark_fill = PatternFill(start_color="2D2D2D", end_color="2D2D2D", fill_type="solid")
    light_fill = PatternFill(start_color="F7F5F2", end_color="F7F5F2", fill_type="solid")
    white_font = Font(name="Century Gothic", color="FFFFFF", bold=True, size=10)
    normal_font = Font(name="Century Gothic", color="2D2D2D", size=10)
    gold_font = Font(name="Century Gothic", color="B08D57", size=12, bold=True)
    thin_border = Border(
        left=Side(style="thin", color="DDDDDD"),
        right=Side(style="thin", color="DDDDDD"),
        top=Side(style="thin", color="DDDDDD"),
        bottom=Side(style="thin", color="DDDDDD"),
    )

    if os.path.exists(excel_path):
        wb = load_workbook(excel_path)
    else:
        wb = Workbook()

    # In2026 sheet
    if "In2026" in wb.sheetnames:
        ws_in = wb["In2026"]
    else:
        ws_in = wb.active
        ws_in.title = "In2026"
        headers = ["N Factura", "Cliente", "Fecha", "Mes", "Hecha", "Enviada", "Total", "Estado", "Concepto"]
        for ci, h in enumerate(headers, 1):
            cell = ws_in.cell(row=1, column=ci, value=h)
            cell.font = white_font
            cell.fill = dark_fill
            cell.alignment = Alignment(horizontal="center")
            cell.border = thin_border
        widths = [14, 28, 12, 12, 10, 10, 14, 12, 35]
        for i, w in enumerate(widths, 1):
            ws_in.column_dimensions[chr(64 + i)].width = w

    mes_str = MESES_ES.get(fecha.month, "").upper()
    new_row = ws_in.max_row + 1
    in2026_values = [
        num_doc, pagador_nombre, fecha.strftime("%d/%m/%Y"), mes_str,
        "OK", "no", float(total), estado,
        "{} -- {}".format(metodo.upper() if metodo else "", alumno_nombre),
    ]
    for ci, val in enumerate(in2026_values, 1):
        cell = ws_in.cell(row=new_row, column=ci, value=val)
        cell.font = normal_font
        cell.border = thin_border
        if new_row % 2 == 0:
            cell.fill = light_fill
        if ci == 7:
            cell.number_format = "#,##0.00 EUR"

    # Ingresos sheet (Taxfix) - only for facturas
    if es_factura_fiscal:
        if "Ingresos" in wb.sheetnames:
            ws_ing = wb["Ingresos"]
        else:
            ws_ing = wb.create_sheet("Ingresos")
            ing_headers = [
                "NIF - Cliente", "Codigo pais", "Nombre/Razon social",
                "N Factura", "Fecha factura", "Fecha devengo",
                "Tipo factura", "Concepto", "Base Imponible (EUR)", "Total (EUR)"
            ]
            for ci, h in enumerate(ing_headers, 1):
                cell = ws_ing.cell(row=1, column=ci, value=h)
                cell.font = white_font
                cell.fill = dark_fill
                cell.alignment = Alignment(horizontal="center")
                cell.border = thin_border
            for i, w in enumerate([14, 12, 30, 16, 14, 14, 12, 45, 16, 16], 1):
                ws_ing.column_dimensions[chr(64 + i)].width = w

        tipo_fis = tipo_fiscal(pagador_nif)
        concepto_tax = "Clases de inglés. Academia. {}. {}.".format(grupo_nombre, alumno_nombre)
        new_ing_row = ws_ing.max_row + 1
        ing_values = [
            pagador_nif or "", "ES" if pagador_nif else "",
            pagador_nombre, num_doc,
            fecha.strftime("%d/%m/%Y"), fecha.strftime("%d/%m/%Y"),
            tipo_fis, concepto_tax, float(total), float(total),
        ]
        for ci, val in enumerate(ing_values, 1):
            cell = ws_ing.cell(row=new_ing_row, column=ci, value=val)
            cell.font = normal_font
            cell.border = thin_border
            if new_ing_row % 2 == 0:
                cell.fill = light_fill
            if ci in (9, 10):
                cell.number_format = "#,##0.00 EUR"

    wb.save(excel_path)
    return excel_path


def delete_from_excel(num_doc, year=None):
    if year is None:
        year = date.today().year
    excel_path = os.path.join(BASE_INVOICE_DIR, "Facturas {}".format(year), "Registro_Ingresos_{}.xlsx".format(year))
    if not os.path.exists(excel_path):
        return
    from openpyxl import load_workbook
    wb = load_workbook(excel_path)
    for sheet_name in ["In2026", "Ingresos"]:
        if sheet_name not in wb.sheetnames:
            continue
        ws = wb[sheet_name]
        rows_to_delete = []
        for row in ws.iter_rows(min_row=2):
            for cell in row:
                if str(cell.value) == str(num_doc):
                    rows_to_delete.append(cell.row)
                    break
        for row_num in sorted(rows_to_delete, reverse=True):
            ws.delete_rows(row_num)
    wb.save(excel_path)


def generate_invoice_for_pago(pago, tipo="factura"):
    acad = pago.academia
    pagador = pago.pagador
    alumno = pago.alumno
    grupo = pago.grupo
    extras = pago.extras or []
    metodo = pago.metodo or ""

    es_factura_fiscal = es_factura(metodo)
    tipo_doc = "factura" if es_factura_fiscal else "recibo"

    num_doc, file_path, docx_path = generate_invoice(
        academia_nombre=acad.academia_nombre or "Camila Pereiras Casal",
        academia_dir=acad.academia_dir or "C/ Pedro Soto Couselo 5, 2 B",
        academia_ciudad="36995, Poio (Pontevedra)",
        academia_tel=acad.academia_tel or "698 183 419",
        academia_cif=acad.academia_nif or "39468659S",
        pagador_nombre=pagador.nombre,
        pagador_nif=getattr(pagador, "nif", "") or "",
        pagador_dir="",
        pagador_ciudad="",
        pagador_tel=getattr(pagador, "telefono", "") or "",
        pagador_email=getattr(pagador, "email", "") or "",
        alumno_nombre=alumno.nombre,
        grupo_nombre=grupo.nombre if grupo else "",
        periodo=pago.periodo,
        mensualidad=pago.mensualidad,
        descuento=pago.descuento,
        extras=extras,
        total=pago.total,
        metodo=metodo,
        concepto_libre=getattr(pago, "concepto_libre", "") or "",
        num_doc=None,
        fecha=pago.fecha,
        tipo=tipo_doc,
        academia_id=acad.id,
    )

    try:
        extras_total = sum(float(ex.get("importe", 0)) for ex in extras)
        update_excel_log(
            num_doc=num_doc,
            fecha=pago.fecha or date.today(),
            pagador_nombre=pagador.nombre,
            pagador_nif=getattr(pagador, "nif", "") or "",
            pagador_dir="",
            alumno_nombre=alumno.nombre,
            grupo_nombre=grupo.nombre if grupo else "",
            periodo=pago.periodo,
            mensualidad=pago.mensualidad,
            descuento=pago.descuento,
            extras_total=extras_total,
            total=pago.total,
            metodo=metodo,
            estado=pago.estado,
            es_factura_fiscal=es_factura_fiscal,
        )
    except Exception as e:
        print("Excel log error (non-critical): {}".format(e))

    return num_doc, file_path
